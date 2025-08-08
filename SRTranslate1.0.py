import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox, filedialog
import json
import requests
import threading
import os
from openai import OpenAI
from docx import Document  # 用于处理Word文档
import io
import unicodedata  # 用于处理特殊字符


class DeepSeekTranslator:
    def __init__(self, root):
        self.root = root
        self.root.title("DeepSeek 翻译器 - 支持Word文档")
        self.root.geometry("1000x700")

        # 自定义词典
        self.custom_dict = {}

        # 创建导航栏和主界面
        self.create_navigation()

        # 加载配置文件
        self.load_config()

    def create_navigation(self):
        # 创建导航栏
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill=tk.BOTH, expand=True)

        # 创建三个标签页
        self.create_translation_tab()
        self.create_dictionary_tab()
        self.create_api_settings_tab()

    def create_translation_tab(self):
        # 首页 - 翻译标签页
        self.translation_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.translation_tab, text="首页")

        # 使用PanedWindow分隔输入和输出区域
        main_paned = tk.PanedWindow(self.translation_tab, orient=tk.HORIZONTAL, sashrelief=tk.RAISED, sashwidth=5)
        main_paned.pack(fill=tk.BOTH, expand=True)

        # 输入区域
        input_frame = ttk.Frame(main_paned)
        main_paned.add(input_frame)

        # 输出区域
        output_frame = ttk.Frame(main_paned)
        main_paned.add(output_frame)

        # 输入区域内容
        input_label = ttk.Label(input_frame, text="输入文本", font=('Arial', 10, 'bold'))
        input_label.pack(pady=(5, 0))

        self.input_text = scrolledtext.ScrolledText(input_frame, wrap=tk.WORD, font=('Arial', 10))
        self.input_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # 输入按钮区域
        input_button_frame = ttk.Frame(input_frame)
        input_button_frame.pack(fill=tk.X, padx=5, pady=5)

        translate_button = ttk.Button(input_button_frame, text="翻译", command=self.start_translation)
        translate_button.pack(side=tk.LEFT, padx=5)

        clear_button = ttk.Button(input_button_frame, text="清空", command=self.clear_text)
        clear_button.pack(side=tk.LEFT, padx=5)

        load_file_button = ttk.Button(input_button_frame, text="加载文本文件", command=self.load_text_file)
        load_file_button.pack(side=tk.LEFT, padx=5)

        load_word_button = ttk.Button(input_button_frame, text="加载Word文档", command=self.load_word_file)
        load_word_button.pack(side=tk.LEFT, padx=5)

        # 输出区域内容
        output_label = ttk.Label(output_frame, text="翻译结果", font=('Arial', 10, 'bold'))
        output_label.pack(pady=(5, 0))

        # 使用Text控件而不是ScrolledText以便更好地处理emoji
        self.result_text = tk.Text(output_frame, wrap=tk.WORD, state=tk.DISABLED,
                                   font=('Arial', 10), height=20)
        scrollbar = ttk.Scrollbar(output_frame, command=self.result_text.yview)
        self.result_text.configure(yscrollcommand=scrollbar.set)

        self.result_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5, pady=5)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # 输出按钮区域
        output_button_frame = ttk.Frame(output_frame)
        output_button_frame.pack(fill=tk.X, padx=5, pady=5)

        copy_button = ttk.Button(output_button_frame, text="复制结果", command=self.copy_result)
        copy_button.pack(side=tk.LEFT, padx=5)

        clear_result_button = ttk.Button(output_button_frame, text="清空结果", command=self.clear_result)
        clear_result_button.pack(side=tk.LEFT, padx=5)

        save_result_button = ttk.Button(output_button_frame, text="保存为文本", command=self.save_as_text)
        save_result_button.pack(side=tk.LEFT, padx=5)

        save_word_button = ttk.Button(output_button_frame, text="保存为Word", command=self.save_as_word)
        save_word_button.pack(side=tk.LEFT, padx=5)

    def create_dictionary_tab(self):
        # 词典标签页
        self.dictionary_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.dictionary_tab, text="词典")

        # 词典操作区域
        dict_operation_frame = ttk.LabelFrame(self.dictionary_tab, text="词典操作")
        dict_operation_frame.pack(fill=tk.X, padx=10, pady=10)

        # 添加词汇区域
        add_frame = ttk.Frame(dict_operation_frame)
        add_frame.pack(fill=tk.X, pady=5)

        ttk.Label(add_frame, text="原词:").pack(side=tk.LEFT, padx=5)
        self.original_word = ttk.Entry(add_frame, width=25)
        self.original_word.pack(side=tk.LEFT, padx=5)

        ttk.Label(add_frame, text="翻译:").pack(side=tk.LEFT, padx=5)
        self.translated_word = ttk.Entry(add_frame, width=25)
        self.translated_word.pack(side=tk.LEFT, padx=5)

        add_button = ttk.Button(add_frame, text="添加", command=self.add_custom_word)
        add_button.pack(side=tk.LEFT, padx=5)

        # 词典文件操作按钮
        file_button_frame = ttk.Frame(dict_operation_frame)
        file_button_frame.pack(fill=tk.X, pady=5)

        load_dict_button = ttk.Button(file_button_frame, text="加载词典", command=self.load_dict)
        load_dict_button.pack(side=tk.LEFT, padx=5)

        save_dict_button = ttk.Button(file_button_frame, text="保存词典", command=self.save_dict)
        save_dict_button.pack(side=tk.LEFT, padx=5)

        # 词典显示区域
        dict_display_frame = ttk.LabelFrame(self.dictionary_tab, text="自定义词典内容")
        dict_display_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=(0, 10))

        self.dict_tree = ttk.Treeview(dict_display_frame, columns=('original', 'translated'), show='headings')
        self.dict_tree.heading('original', text='原词')
        self.dict_tree.heading('translated', text='翻译')
        self.dict_tree.column('original', width=200)
        self.dict_tree.column('translated', width=200)

        vsb = ttk.Scrollbar(dict_display_frame, orient="vertical", command=self.dict_tree.yview)
        hsb = ttk.Scrollbar(dict_display_frame, orient="horizontal", command=self.dict_tree.xview)
        self.dict_tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)

        self.dict_tree.grid(row=0, column=0, sticky='nsew')
        vsb.grid(row=0, column=1, sticky='ns')
        hsb.grid(row=1, column=0, sticky='ew')

        dict_display_frame.grid_rowconfigure(0, weight=1)
        dict_display_frame.grid_columnconfigure(0, weight=1)

        # 右键菜单
        self.dict_menu = tk.Menu(self.root, tearoff=0)
        self.dict_menu.add_command(label="删除", command=self.delete_selected_word)
        self.dict_tree.bind("<Button-3>", self.show_dict_context_menu)

    def create_api_settings_tab(self):
        # API设置标签页
        self.api_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.api_tab, text="API设置")

        # API设置区域
        api_frame = ttk.LabelFrame(self.api_tab, text="DeepSeek API 设置")
        api_frame.pack(fill=tk.BOTH, padx=10, pady=10)

        ttk.Label(api_frame, text="API Key:").grid(row=0, column=0, padx=10, pady=10, sticky=tk.W)
        self.api_key_entry = ttk.Entry(api_frame, width=50)
        self.api_key_entry.grid(row=0, column=1, padx=10, pady=10, sticky=tk.W)

        # 语言选择
        lang_frame = ttk.LabelFrame(api_frame, text="默认翻译设置")
        lang_frame.grid(row=1, column=0, columnspan=2, padx=10, pady=10, sticky=tk.W)

        ttk.Label(lang_frame, text="源语言:").grid(row=0, column=0, padx=5, pady=5, sticky=tk.W)
        self.source_lang = ttk.Combobox(lang_frame, values=["auto", "zh", "en", "ja", "ko", "fr", "de", "es", "ru"])
        self.source_lang.grid(row=0, column=1, padx=5, pady=5, sticky=tk.W)
        self.source_lang.set("auto")

        ttk.Label(lang_frame, text="目标语言:").grid(row=0, column=2, padx=5, pady=5, sticky=tk.W)
        self.target_lang = ttk.Combobox(lang_frame, values=["zh", "en", "ja", "ko", "fr", "de", "es", "ru"])
        self.target_lang.grid(row=0, column=3, padx=5, pady=5, sticky=tk.W)
        self.target_lang.set("zh")

        # 保存设置按钮
        save_button = ttk.Button(api_frame, text="保存设置", command=self.save_config)
        save_button.grid(row=2, column=0, columnspan=2, pady=10)

    def add_custom_word(self):
        original = self.original_word.get().strip()
        translated = self.translated_word.get().strip()

        if original and translated:
            self.custom_dict[original] = translated
            self.update_dict_tree()
            messagebox.showinfo("成功", f"已添加自定义翻译: {original} -> {translated}")
            self.original_word.delete(0, tk.END)
            self.translated_word.delete(0, tk.END)
        else:
            messagebox.showwarning("警告", "请输入原词和翻译")

    def update_dict_tree(self):
        self.dict_tree.delete(*self.dict_tree.get_children())
        for original, translated in self.custom_dict.items():
            self.dict_tree.insert('', tk.END, values=(original, translated))

    def delete_selected_word(self):
        selected_item = self.dict_tree.selection()
        if selected_item:
            original = self.dict_tree.item(selected_item)['values'][0]
            del self.custom_dict[original]
            self.dict_tree.delete(selected_item)
            messagebox.showinfo("成功", f"已删除: {original}")

    def show_dict_context_menu(self, event):
        item = self.dict_tree.identify_row(event.y)
        if item:
            self.dict_tree.selection_set(item)
            self.dict_menu.post(event.x_root, event.y_root)

    def start_translation(self):
        api_key = self.api_key_entry.get().strip()
        if not api_key:
            messagebox.showerror("错误", "请输入 API Key")
            return

        text = self.input_text.get("1.0", tk.END).strip()
        if not text:
            messagebox.showerror("错误", "请输入要翻译的文本")
            return

        # 显示进度条
        self.show_progress()

        # 在新线程中执行翻译
        threading.Thread(
            target=self.translate_text,
            args=(text, api_key, self.source_lang.get(), self.target_lang.get()),
            daemon=True
        ).start()

    def translate_text(self, text, api_key, source_lang, target_lang):
        # 先应用自定义词典
        for original, translated in self.custom_dict.items():
            text = text.replace(original, translated)

        try:
            # 使用 OpenAI SDK 调用 DeepSeek API
            client = OpenAI(
                api_key=api_key,
                base_url="https://api.deepseek.com"
            )

            # 构建翻译提示词，明确要求保留emoji和特殊字符
            translation_prompt = (
                f"请将以下{source_lang}内容翻译成{target_lang}，"
                f"保持专业准确的翻译风格。"
                f"请保留所有emoji表情符号和特殊字符，不要修改或删除它们:\n\n{text}"
            )

            response = client.chat.completions.create(
                model="deepseek-chat",
                messages=[
                    {"role": "system", "content": "你是一名专业的翻译助手，擅长保留原文中的emoji和特殊符号"},
                    {"role": "user", "content": translation_prompt},
                ],
                stream=False
            )

            result = response.choices[0].message.content

            # 确保特殊字符和emoji被正确处理
            result = self.normalize_text(result)
            self.root.after(0, self.display_result, result)

        except Exception as e:
            error_msg = f"翻译出错: {str(e)}"
            # 如果是 API 错误，尝试获取更详细的错误信息
            if hasattr(e, 'response') and e.response:
                try:
                    error_detail = e.response.json().get('error', {}).get('message', '')
                    error_msg += f"\n详细信息: {error_detail}"
                except:
                    pass
            self.root.after(0, messagebox.showerror, "翻译错误", error_msg)
        finally:
            self.root.after(0, self.hide_progress)

    def normalize_text(self, text):
        """规范化文本，确保特殊字符和emoji正确显示"""
        # 使用unicodedata规范化文本
        normalized = unicodedata.normalize('NFKC', text)
        return normalized

    def show_progress(self):
        if not hasattr(self, 'progress_frame'):
            self.progress_frame = ttk.Frame(self.translation_tab)
            self.progress_frame.pack(fill=tk.X, padx=5, pady=5)

            self.progress_label = ttk.Label(self.progress_frame, text="正在翻译...")
            self.progress_label.pack(side=tk.LEFT, padx=5)

            self.progress_bar = ttk.Progressbar(self.progress_frame, mode="indeterminate")
            self.progress_bar.pack(fill=tk.X, expand=True, padx=5)

        self.progress_bar.start()
        self.progress_frame.pack(fill=tk.X, padx=5, pady=5)

    def hide_progress(self):
        if hasattr(self, 'progress_frame'):
            self.progress_bar.stop()
            self.progress_frame.pack_forget()

    def display_result(self, result):
        self.result_text.config(state=tk.NORMAL)
        self.result_text.delete("1.0", tk.END)

        # 确保使用能显示emoji的字体
        self.result_text.configure(font=('Segoe UI Emoji', 10))

        # 插入文本，确保emoji和特殊字符正确显示
        self.result_text.insert(tk.END, result)
        self.result_text.config(state=tk.DISABLED)

    def copy_result(self):
        result = self.result_text.get("1.0", tk.END).strip()
        if result:
            try:
                # 使用clipboard_clear和clipboard_append来复制包含emoji的文本
                self.root.clipboard_clear()
                self.root.clipboard_append(result)
                messagebox.showinfo("成功", "已复制翻译结果到剪贴板")
            except Exception as e:
                messagebox.showerror("错误", f"复制失败: {str(e)}")

    def clear_result(self):
        self.result_text.config(state=tk.NORMAL)
        self.result_text.delete("1.0", tk.END)
        self.result_text.config(state=tk.DISABLED)

    def save_as_text(self):
        result = self.result_text.get("1.0", tk.END).strip()
        if not result:
            messagebox.showwarning("警告", "没有可保存的翻译结果")
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Text Files", "*.txt"), ("All Files", "*.*")]
        )

        if file_path:
            try:
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(result)
                messagebox.showinfo("成功", "翻译结果已保存为文本文件")
            except Exception as e:
                messagebox.showerror("错误", f"无法保存文件: {str(e)}")

    def save_as_word(self):
        result = self.result_text.get("1.0", tk.END).strip()
        if not result:
            messagebox.showwarning("警告", "没有可保存的翻译结果")
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )

        if file_path:
            try:
                doc = Document()
                doc.add_paragraph(result)
                doc.save(file_path)
                messagebox.showinfo("成功", "翻译结果已保存为Word文档")
            except Exception as e:
                messagebox.showerror("错误", f"无法保存Word文档: {str(e)}")

    def clear_text(self):
        self.input_text.delete("1.0", tk.END)

    def load_text_file(self):
        file_path = filedialog.askopenfilename(
            filetypes=[("Text Files", "*.txt"), ("All Files", "*.*")]
        )

        if file_path:
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
                self.input_text.delete("1.0", tk.END)
                self.input_text.insert(tk.END, content)
            except Exception as e:
                messagebox.showerror("错误", f"无法加载文本文件: {str(e)}")

    def load_word_file(self):
        file_path = filedialog.askopenfilename(
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )

        if file_path:
            try:
                doc = Document(file_path)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)

                content = "\n".join(full_text)
                self.input_text.delete("1.0", tk.END)
                self.input_text.insert(tk.END, content)
            except Exception as e:
                messagebox.showerror("错误", f"无法加载Word文档: {str(e)}")

    def save_dict(self):
        file_path = filedialog.asksaveasfilename(
            defaultextension=".json",
            filetypes=[("JSON Files", "*.json"), ("All Files", "*.*")]
        )

        if file_path:
            try:
                with open(file_path, "w", encoding="utf-8") as f:
                    json.dump(self.custom_dict, f, ensure_ascii=False, indent=2)
                messagebox.showinfo("成功", "自定义词典已保存")
            except Exception as e:
                messagebox.showerror("错误", f"无法保存词典: {str(e)}")

    def load_dict(self):
        file_path = filedialog.askopenfilename(
            filetypes=[("JSON Files", "*.json"), ("All Files", "*.*")]
        )

        if file_path:
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    self.custom_dict = json.load(f)
                self.update_dict_tree()
                messagebox.showinfo("成功", "自定义词典已加载")
            except Exception as e:
                messagebox.showerror("错误", f"无法加载词典: {str(e)}")

    def load_config(self):
        config_path = "translator_config.json"
        if os.path.exists(config_path):
            try:
                with open(config_path, "r", encoding="utf-8") as f:
                    config = json.load(f)
                    self.api_key_entry.insert(0, config.get("api_key", ""))
                    self.source_lang.set(config.get("source_lang", "auto"))
                    self.target_lang.set(config.get("target_lang", "zh"))

                    # 加载自定义词典
                    if "custom_dict" in config:
                        self.custom_dict = config["custom_dict"]
                        self.update_dict_tree()
            except Exception as e:
                print(f"加载配置出错: {str(e)}")

    def save_config(self):
        config = {
            "api_key": self.api_key_entry.get(),
            "source_lang": self.source_lang.get(),
            "target_lang": self.target_lang.get(),
            "custom_dict": self.custom_dict
        }

        try:
            with open("translator_config.json", "w", encoding="utf-8") as f:
                json.dump(config, f, indent=2)
            messagebox.showinfo("成功", "配置已保存")
        except Exception as e:
            messagebox.showerror("错误", f"无法保存配置: {str(e)}")

    def on_closing(self):
        self.save_config()
        self.root.destroy()


if __name__ == "__main__":
    root = tk.Tk()

    # 设置能更好显示emoji的字体
    try:
        # Windows系统
        root.option_add("*Font", ("Segoe UI Emoji", 10))
    except:
        try:
            # macOS系统
            root.option_add("*Font", ("Apple Color Emoji", 12))
        except:
            # 其他系统使用默认字体
            pass

    app = DeepSeekTranslator(root)
    root.protocol("WM_DELETE_WINDOW", app.on_closing)
    root.mainloop()